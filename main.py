import os
import subprocess
import datetime as dt
import tkinter as tk
from tkinter import filedialog, messagebox
import docx 
from docx2pdf import convert

class InvoiceAutomation:
    def __init__(self):
        
        self.root = tk.Tk()
        self.root.title("Invoice Automation")
        self.root.geometry("500x600")

        
        self.partner_label = tk.Label(self.root, text="Partner")
        self.partner_street_label = tk.Label(self.root, text="Street")
        self.partner_zip_city_country_label = tk.Label(self.root, text="ZIP / City / Country")
        self.invoice_number_label = tk.Label(self.root, text="Invoice Number")
        self.service_description_label = tk.Label(self.root, text="Service Description")
        self.service_amount_label = tk.Label(self.root, text="Service Amount")
        self.service_single_price_label = tk.Label(self.root, text="Service Single Price")
        self.payment_method_label = tk.Label(self.root, text="Payment Method")

        
        self.payment_methods = {
            "Main Bank": {
                "recipient": "Happy Company",
                "bank": "Hello World Bank",
                "iban": "XY12 3456 7890 1234",
                "bic": "ABCDEFGH"
            },
            "Second Bank": {
                "recipient": "Happy Company",
                "bank": "Second World Bank",
                "iban": "XY98 7654 3210 9876",
                "bic": "IJKLMNOP"
            },
            "Private Bank": {
                "recipient": "Happy Company",
                "bank": "Private Bank",
                "iban": "XY11 2233 4455 6677",
                "bic": "QRSTUVWX"
            }
        }

       
        self.partner_entry = tk.Entry(self.root)
        self.partner_street_entry = tk.Entry(self.root)
        self.partner_zip_city_country_entry = tk.Entry(self.root)
        self.invoice_number_entry = tk.Entry(self.root)
        self.service_description_entry = tk.Entry(self.root)
        self.service_amount_entry = tk.Entry(self.root)
        self.service_single_price_entry = tk.Entry(self.root)

        
        self.payment_method = tk.StringVar(self.root)
        self.payment_method.set("Main Bank")
        self.payment_method_dropdown = tk.OptionMenu(
            self.root,
            self.payment_method,
            "Main Bank",
            "Second Bank",
            "Private Bank"
        )

        
        self.create_button = tk.Button(
            self.root,
            text="Create Invoice",
            command=self.create_invoice
        )

        
        padding_options = {"fill": "x", "expand": True, "padx": 5, "pady": 2}

        self.partner_label.pack(**padding_options)
        self.partner_entry.pack(**padding_options)

        self.partner_street_label.pack(**padding_options)
        self.partner_street_entry.pack(**padding_options)

        self.partner_zip_city_country_label.pack(**padding_options)
        self.partner_zip_city_country_entry.pack(**padding_options)

        self.invoice_number_label.pack(**padding_options)
        self.invoice_number_entry.pack(**padding_options)

        self.service_description_label.pack(**padding_options)
        self.service_description_entry.pack(**padding_options)

        self.service_amount_label.pack(**padding_options)
        self.service_amount_entry.pack(**padding_options)

        self.service_single_price_label.pack(**padding_options)
        self.service_single_price_entry.pack(**padding_options)

        self.payment_method_label.pack(**padding_options)
        self.payment_method_dropdown.pack(**padding_options)

        self.create_button.pack(**padding_options)

        self.root.mainloop()

    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    def create_invoice(self):
        
        doc = docx.Document("template.docx")

        
        selected_payment = self.payment_methods[self.payment_method.get()]

        
        try:
            replacements = {
                "[DATE]": dt.datetime.today().strftime("%Y-%m-%d"),
                "[PARTNER]": self.partner_entry.get(),
                "[PARTNER_STREET]": self.partner_street_entry.get(),
                "[PARTNER_ZIP_CITY_COUNTRY]": self.partner_zip_city_country_entry.get(),
                "[INVOICE_NUMBER]": self.invoice_number_entry.get(),
                "[SERVICE_DESCRIPTION]": self.service_description_entry.get(),
                "[AMOUNT]": self.service_amount_entry.get(),
                "[SINGLE_PRICE]": f"${float(self.service_single_price_entry.get()):.2f}",
                "[FULL_PRICE]": f"${float(self.service_amount_entry.get()) * float(self.service_single_price_entry.get()):.2f}",
                "[RECIPIENT]": selected_payment["recipient"],
                "[BANK]": selected_payment["bank"],
                "[IBAN]": selected_payment["iban"],
                "[BIC]": selected_payment["bic"],
            }
        except ValueError:
            messagebox.showerror("Error", "Invalid amount or price")
            return

        
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)

        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)

        
        save_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Documents", "*.pdf")]
        )
        if not save_path:
            return

        
        temp_docx = "filled.docx"
        temp_pdf = "filled.pdf"
        doc.save(temp_docx)

       
        try:
           
             convert(temp_docx, save_path)
            
            
        except Exception as e:
            
            messagebox.showerror("Error", f"Failed to convert to PDF: {e}")
            return
        finally:
           
            if os.path.exists(temp_docx):
                try:
                    os.remove(temp_docx)
                except Exception:
                    pass

        messagebox.showinfo("Success", "Invoice created and saved successfully")


if __name__ == "__main__":
    InvoiceAutomation()
